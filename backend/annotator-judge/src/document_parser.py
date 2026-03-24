# src/document_parser.py
import os
from pathlib import Path


def read_document(file_path: str) -> str:
    """
    Reads and extracts text from a document.
    Supports .pdf, .docx, and .txt files.

    Args:
        file_path: Path to the instruction document.

    Returns:
        Extracted text content as a string.

    Raises:
        FileNotFoundError: If the file doesn't exist.
        NotImplementedError: If the file type is not supported.
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"Document not found: {file_path}")

    extension = path.suffix.lower()

    if extension == ".txt":
        return _read_txt(path)
    elif extension == ".pdf":
        return _read_pdf(path)
    elif extension == ".docx":
        return _read_docx(path)
    else:
        raise NotImplementedError(
            f"File type '{extension}' is not supported. Use .pdf, .docx, or .txt"
        )


def _read_txt(path: Path) -> str:
    """Read plain text file."""
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def _read_pdf(path: Path) -> str:
    """Read PDF file using PyMuPDF (fitz) with table extraction."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ImportError(
            "PyMuPDF is required to read PDF files. Install with: pip install PyMuPDF"
        )

    doc = fitz.open(path)
    text_parts = []

    for page in doc:
        # Try to extract tables first (PyMuPDF 1.23.0+)
        try:
            tables = page.find_tables()
            if tables and len(tables.tables) > 0:
                # Page has tables - extract them as markdown
                page_text_parts = []

                # Get regular text blocks
                blocks = page.get_text("blocks")
                table_rects = [t.bbox for t in tables.tables]

                for block in blocks:
                    block_rect = fitz.Rect(block[:4])
                    # Check if block overlaps with any table
                    in_table = any(
                        block_rect.intersects(fitz.Rect(tr)) for tr in table_rects
                    )
                    if not in_table and block[4].strip():
                        page_text_parts.append(block[4].strip())

                # Add tables as markdown
                for table in tables.tables:
                    md_table = _convert_pdf_table_to_markdown(table)
                    if md_table:
                        page_text_parts.append(md_table)

                text_parts.append("\n\n".join(page_text_parts))
            else:
                # No tables - regular extraction
                text_parts.append(page.get_text())
        except AttributeError:
            # Older PyMuPDF version - fall back to regular extraction
            text_parts.append(page.get_text())

    doc.close()
    return "\n\n".join(text_parts)


def _convert_pdf_table_to_markdown(table) -> str:
    """Convert PyMuPDF table to markdown format."""
    try:
        data = table.extract()
        if not data or len(data) == 0:
            return ""

        rows_text = []
        for i, row in enumerate(data):
            cells = [
                str(cell).strip().replace("\n", " ") if cell else "" for cell in row
            ]
            row_text = " | ".join(cells)
            rows_text.append(f"| {row_text} |")

            # Add separator after header row
            if i == 0:
                separator = "|" + "|".join(["---"] * len(cells)) + "|"
                rows_text.append(separator)

        return "\n".join(rows_text)
    except Exception:
        return ""


def _read_docx(path: Path) -> str:
    """Read Word document using python-docx, including tables and text boxes."""
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError:
        raise ImportError(
            "python-docx is required to read DOCX files. Install with: pip install python-docx"
        )

    doc = Document(path)
    text_parts = []

    try:
        # Iterate through body elements to preserve order (paragraphs + tables)
        for element in doc.element.body:
            # Paragraph
            if element.tag.endswith("p"):
                for para in doc.paragraphs:
                    if para._element is element:
                        if para.text.strip():
                            text_parts.append(para.text)
                        # Also check for text boxes within paragraphs
                        textbox_content = _extract_textboxes_from_element(element)
                        if textbox_content:
                            text_parts.append(textbox_content)
                        break

            # Table
            elif element.tag.endswith("tbl"):
                for table in doc.tables:
                    if table._tbl is element:
                        table_text = _extract_table_as_text(table)
                        if table_text.strip():
                            text_parts.append(table_text)
                        break
    except Exception as e:
        # Fallback: simple extraction if structured parsing fails
        print(f"  Warning: Structured parsing failed ({e}), using fallback...")
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        for table in doc.tables:
            table_text = _extract_table_as_text(table)
            if table_text.strip():
                text_parts.append(table_text)

    return "\n\n".join(text_parts)


def _extract_textboxes_from_element(element) -> str:
    """Extract text from text boxes within an element."""
    try:
        from docx.oxml.ns import qn

        textbox_texts = []
        # Look for text boxes (w:txbxContent)
        for txbx in element.iter(qn("w:txbxContent")):
            for p in txbx.iter(qn("w:p")):
                text = "".join(node.text for node in p.iter(qn("w:t")) if node.text)
                if text.strip():
                    textbox_texts.append(text.strip())

        return "\n".join(textbox_texts)
    except Exception:
        return ""


def _extract_table_as_text(table) -> str:
    """Extract table content as markdown-style text with better handling."""
    try:
        rows_text = []
        prev_row_cells = None  # Track for merged cell detection

        for i, row in enumerate(table.rows):
            cells = []
            for j, cell in enumerate(row.cells):
                # Get cell text, handling merged cells
                cell_text = cell.text.strip().replace("\n", " ")

                # Check for vertically merged cells (avoid duplicates)
                if prev_row_cells and j < len(prev_row_cells):
                    if cell._tc is prev_row_cells[j]._tc:
                        cell_text = ""  # Skip repeated merged cell

                cells.append(cell_text)

            row_text = " | ".join(cells)
            rows_text.append(f"| {row_text} |")

            # Add separator after header row
            if i == 0:
                separator = "|" + "|".join(["---"] * len(cells)) + "|"
                rows_text.append(separator)

            prev_row_cells = row.cells

        return "\n".join(rows_text)
    except Exception as e:
        # Fallback: simple extraction
        try:
            rows_text = []
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                rows_text.append("| " + " | ".join(cells) + " |")
            return "\n".join(rows_text)
        except Exception:
            return ""
